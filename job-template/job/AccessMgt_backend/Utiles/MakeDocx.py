import uuid
from docx import Document
from docx.shared import RGBColor
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
import matplotlib.pyplot as plt
import io
from Utiles import global_var


class MakeDocxClass():
    """
    用于生成word文档
    """

    @staticmethod
    def make_docx(system_info,basic_info,tree_data,tree_details,tree_results):
        """
        生成word
        :param info:
        :return:
        """
        normalize_func_dict = {
            "gaussian_mapping":"高斯函数映射",
            "cauchy_mapping":"柯西分布",
            "exponential_decay":"指数衰减函数",
            "logistic_variant":"辑斯蒂函数变体",
            "linear_segment":" 对称线性函数",
            "linear_mapping":"单段线性函数",
            "zero_one_mapping":"开关函数",
            "segmentation_liner_mapping":"线性分段函数",
            "segmentation_mapping":"台阶分段函数",
        }
        selected_algorithm_dict= {
            "weighted":"加权平均法",
            "grey":"灰色白化权聚类法",
            "fuzzy":"模糊综合评判法",
            "index":"幂指数",
            "expert":"专家打分法",
        }
        selected_analysis_method_dict = {
            "ahp":"层次分析法",
            "CoefficientMethod":"变异系数法",
            "PcaWeight":"主成分分析法",
            "FactorAnalysis":"因子分析法",
            "EntropyWeight":"熵权法",
            "GreyRelationalAnalysis":"灰色关联分析法",
            "DareMethod":"环比系数法",
            "directEmpowerment":"直接赋权法",
        }
        # 创建word对象
        doc = Document()
        # 插入标题
        doc = MakeDocxClass.set_title("作战效能分析报告",doc)
        # 插入一级标题
        doc = MakeDocxClass.add_heading1("评估体系",doc)
        # 插入正文
        doc = MakeDocxClass.set_text("体系名称:{}".format(system_info["name"]),doc)
        doc = MakeDocxClass.set_text("体系创建时间:{}".format(system_info["createTime"]),doc)
        doc = MakeDocxClass.set_text("体系描述:{}".format(system_info["description"]),doc)
        # 插入一级标题
        doc = MakeDocxClass.add_heading1("场景信息",doc)
        # 插入正文
        doc = MakeDocxClass.set_text("场景名称:{}".format(basic_info["name"]), doc)
        doc = MakeDocxClass.set_text("场景数据录入时间:{}".format(basic_info["createTime"]), doc)
        doc = MakeDocxClass.set_text("场景描述:{}".format(basic_info["description"]), doc)
        # 插入一级标题
        doc = MakeDocxClass.add_heading1("评估结果",doc)
        # 插入二级标题
        doc = MakeDocxClass.add_heading2("总体得分",doc)
        # 插入正文总体得分
        doc = MakeDocxClass.set_text("总得分：{}".format(tree_results[tree_data["id"]]),doc)
        # 插入二级标题
        doc = MakeDocxClass.add_heading2("各项指标得分",doc)
        # 插入表格
        doc = MakeDocxClass.insert_hierarchical_table(doc,[tree_data],tree_details,tree_results)
        # 插入二级标题
        doc = MakeDocxClass.add_heading2("各项指标设置", doc)
        # 节点详情插入各个节点详细数据
        node_id_dict = MakeDocxClass.check_node_children([tree_data])
        print(node_id_dict)
        for node_id,value in node_id_dict.items():
            doc = MakeDocxClass.set_text("指标名称:{}".format(tree_details[node_id]["label"]),doc)
            node_detail = tree_details[node_id]
            if value:  # 表示有子节点  要列出指标评估算法、子指标权重、
                doc = MakeDocxClass.set_text("权重分析方法:{}".format(selected_analysis_method_dict[node_detail["selectedAnalysisMethod"]]), doc)
                doc = MakeDocxClass.set_text("指标算法:{}".format(selected_algorithm_dict[node_detail["selectedAlgorithm"]]), doc)

            else:  # 表示没有子节点
                if node_detail['scoringMethod'] == "expert":
                    doc = MakeDocxClass.set_text("评分方式:{} , 打分结果:{}".format("专家打分法",node_detail["expertScore"]), doc)
                else:
                    doc = MakeDocxClass.set_text("评分方式:{} , 公式:{}".format("公式计算", node_detail["formula"]), doc)
                    doc = MakeDocxClass.set_text("归一化算法:{}".format(normalize_func_dict[node_detail["normalizationAlgorithm"]]), doc)
            doc = MakeDocxClass.set_text("", doc)


        # # 插入二级标题
        # doc = MakeDocxClass.add_heading2("二级指标得分情况-饼状图",doc)
        # 获取二级指标列表
        second_node_list = tree_data["children"]
        if second_node_list:
            row_data = []
            labels = []
            doc = MakeDocxClass.add_heading2("二级指标得分情况-柱状图", doc)
            for info in second_node_list:
                node_id = info["id"]
                row_data.append(int(tree_results[node_id]))
                labels.append(tree_details[node_id]['label'])

            img_stream = MakeDocxClass.generate_bar_chart(row_data,labels,"二级指标评分","指标项","分数")
            doc.add_picture(img_stream, width=Inches(6))
        # 生成word
        uid = str(uuid.uuid1())
        print(os.path.join(global_var.get_value("current_dir"),"{}.docx".format(uid)))
        doc.save(os.path.join(global_var.get_value("current_dir"),"WordDir","{}.docx".format(uid)))
        return uid,doc

    @staticmethod
    def generate_bar_chart(data, labels, title, x_label, y_label):
        """生成柱状图并返回图片流"""
        # 设置中文字体，确保中文正常显示
        plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]

        # 创建柱状图
        fig, ax = plt.subplots(figsize=(8, 5))
        bars = ax.bar(labels, data, color=['#4CAF50', '#2196F3', '#FFC107', '#F44336', '#9C27B0'])

        # 添加数据标签
        for bar in bars:
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width() / 2., height + 0.1,
                    f'{height}', ha='center', va='bottom')

        # 设置标题和坐标轴标签
        ax.set_title(title)
        ax.set_xlabel(x_label)
        ax.set_ylabel(y_label)

        # 调整布局
        plt.tight_layout()

        # 将图表保存到字节流
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', bbox_inches='tight', dpi=150)
        img_stream.seek(0)  # 回到流的开始位置

        return img_stream


    @staticmethod
    def check_node_children(tree):
        """
        生成树结构中节点是否有子节点的字典

        参数:
            tree: 树结构，格式为[{"id":..., "children": [...]}, ...]

        返回:
            dict: 键为节点id，值为True（有子节点）或False（无子节点）
        """
        node_children_status = {}

        def traverse(nodes):
            """递归遍历节点及其子节点"""
            for node in nodes:
                node_id = node.get('id')
                if node_id is not None:
                    # 检查是否有子节点（且子节点列表不为空）
                    children = node.get('children', [])
                    has_children = len(children) > 0
                    node_children_status[node_id] = has_children
                # 递归处理子节点
                if children:
                    traverse(children)

        # 开始遍历整个树
        traverse(tree)
        return node_children_status

    @staticmethod
    def insert_hierarchical_table(doc, data, data_details,data_results,header_font_name="微软雅黑", header_font_size=12, content_font_name="宋体",
                                  content_font_size=10):
        """
        向现有文档中插入带字体样式的层级树结构表格

        参数:
        doc: 已打开的Document对象
        data: 层级数据，格式为[{"name":"节点名称", "children":[]}]
        header_font_name: 表头字体名称
        header_font_size: 表头字体大小(磅)
        content_font_name: 内容字体名称
        content_font_size: 内容字体大小(磅)
        """
        # 在表格前添加一个段落作为间隔
        doc.add_paragraph()

        # 提取所有数据并展平为列表
        flat_data = []

        def flatten(structure, level=0, parent=None):
            """递归展平树形结构数据"""
            if isinstance(structure, dict):
                structure = [structure]

            for item in structure:
                new_id = item.get('id',"未找到该指标名称")
                node_name = data_details[new_id]["label"]
                flat_data.append({
                    'name': node_name,
                    'level': level,
                    'parent': parent,
                    'mark': data_results[new_id]
                })

                children = item.get('children', [])
                if children:
                    flatten(children, level + 1, node_name)

        # 展平数据
        flatten(data)

        # 创建表格
        if not flat_data:
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '层级结构'
            hdr_cells[1].text = '说明'
        else:
            table = doc.add_table(rows=len(flat_data) + 1, cols=2)

        # 设置表格列宽
        table.autofit = False
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(4)

        # 设置所有单元格的边框和垂直居中
        for row in table.rows:
            for cell in row.cells:
                MakeDocxClass.set_cell_border(cell, top=4, left=4, bottom=4, right=4, insideH=4, insideV=4)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 设置表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标名称'
        hdr_cells[1].text = '指标得分'

        # 格式化表头（字体、对齐方式）
        for cell in hdr_cells:
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 设置表头字体样式
            MakeDocxClass.set_font_style(para, header_font_name, header_font_size, bold=True)

        # 填充表格内容
        for i, item in enumerate(flat_data, start=1):
            row = table.rows[i]

            # 层级结构列 - 修复多余回车问题
            name_cell = row.cells[0]
            # 清空单元格所有段落
            while len(name_cell.paragraphs) > 0:
                p = name_cell.paragraphs[0]
                name_cell._tc.remove(p._p)
            # 添加新段落并设置内容
            para = name_cell.add_paragraph()
            run = para.add_run(item['name'])  # 直接添加run，不自动创建新段落
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.left_indent = Inches(item['level'] * 0.4)
            # 设置内容字体样式
            MakeDocxClass.set_font_style(para, content_font_name, content_font_size)

            # 说明列 - 修复多余回车问题
            desc_cell = row.cells[1]
            # 清空单元格所有段落
            while len(desc_cell.paragraphs) > 0:
                p = desc_cell.paragraphs[0]
                desc_cell._tc.remove(p._p)
            # 添加新段落并设置内容
            para = desc_cell.add_paragraph()
            # run = para.add_run(f"{item['name']}的说明内容")  # 直接添加run
            # 添加得分
            run = para.add_run("{}".format(int(item["mark"])))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 设置内容字体样式
            MakeDocxClass.set_font_style(para, content_font_name, content_font_size)

        # 在表格后添加一个段落作为间隔
        doc.add_paragraph()

        return doc

    @staticmethod
    def set_title(title_name,doc):
        # 添加一个段落作为主标题
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title_name)

        # 设置字体大小（可以根据需要调整）
        title_run.font.size = Pt(24)  # 24号字体，比默认一级标题更大

        # 设置字体（可选）
        title_run.font.name = "微软雅黑"
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

        # 设置加粗
        title_run.bold = True

        # 设置段落居中对齐
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加一些间距（可选）
        title_paragraph.space_after = Pt(12)
        return doc

    @staticmethod
    def set_text(text,doc):
        """
        设置正文
        :param text:
        :param doc:
        :return:
        """
        # 添加带格式的正文内容
        MakeDocxClass.add_formatted_text(doc, text)
        return doc

    @staticmethod
    def add_heading1(text,doc):
        """插入一级标题并设置为黑色"""
        heading = doc.add_heading(text, level=1)
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)
        return doc

    @staticmethod
    def add_heading2(text,doc):
        """插入二级标题并设置为黑色"""
        heading = doc.add_heading(text, level=2)
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)
        return doc

    @staticmethod
    def add_heading3(text,doc):
        """插入三级标题并设置为黑色"""
        heading = doc.add_heading(text, level=3)
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)
        return doc

    @staticmethod
    def set_cell_border(cell, **kwargs):
        """设置单元格边框样式"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = f'w:{edge}'
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), str(edge_data))  # 1/8磅为单位
                element.set(qn('w:color'), 'auto')

    @staticmethod
    def set_font_style(paragraph, font_name, font_size, bold=False, color=RGBColor(0, 0, 0)):
        """设置段落的字体样式，修复中文字体显示问题"""
        # 确保段落至少有一个run
        if not paragraph.runs:
            paragraph.add_run()

        # 遍历所有run设置字体（处理可能的多个run情况）
        for run in paragraph.runs:
            # 设置西文字体
            run.font.name = font_name
            # 设置字体大小
            run.font.size = Pt(font_size)
            # 设置粗体
            run.font.bold = bold
            # 设置颜色
            run.font.color.rgb = color

            # 解决中文字体显示问题 - 添加东亚字体设置
            r = run._element
            rPr = r.get_or_add_rPr()

            # 设置东亚字体
            if '宋体' in font_name:
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), '宋体')
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rPr.append(rFonts)
            elif '微软雅黑' in font_name:
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), '微软雅黑')
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rPr.append(rFonts)

    @staticmethod
    def set_paragraph_format(paragraph):
        """设置段落格式：确保宋体生效、首行缩进两字符，1.5倍行距"""
        # 设置行距为1.5倍
        paragraph.paragraph_format.line_spacing = 1.5

        # 设置首行缩进两字符（精确计算：中文两字符缩进）
        paragraph.paragraph_format.first_line_indent = Pt(28)  # 微调为28磅更接近实际两字符

        # 设置对齐方式为左对齐
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 设置段前后间距
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)

        # 强制设置字体为宋体（针对中文的完整设置）
        for run in paragraph.runs:
            # 设置西文字体（不影响中文，但需要设置）
            run.font.name = 'Times New Roman'
            # 设置中文字体为宋体
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 设置字体大小
            run.font.size = Pt(12)
            # 添加语言设置
            rPr = run._element.rPr
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run._element.insert(0, rPr)
            lang = OxmlElement('w:lang')
            lang.set(qn('w:val'), 'zh-CN')
            lang.set(qn('w:eastAsia'), 'zh-CN')
            rPr.append(lang)

    @staticmethod
    def add_formatted_text(doc, text):
        """添加带格式的正文内容"""
        paragraph = doc.add_paragraph(text)
        MakeDocxClass.set_paragraph_format(paragraph)
        return paragraph
